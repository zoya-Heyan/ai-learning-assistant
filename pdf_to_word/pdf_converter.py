import os
import logging
import shutil
from pathlib import Path
from typing import Tuple, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

PYPDF2_AVAILABLE = False
pdf2docx_AVAILABLE = False

try:
    from pdf2docx import Converter
    pdf2docx_AVAILABLE = True
    logger.info("pdf2docx is available")
except ImportError:
    logger.warning("pdf2docx not available, install with: pip install pdf2docx")

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
    logger.info("PyPDF2 is available")
except ImportError:
    logger.warning("PyPDF2 not available, install with: pip install PyPDF2")


class PDFToWordConverter:
    """Convert PDF files to editable Word documents"""

    def __init__(self, output_dir: str = "output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.original_dir = self.output_dir / "original"
        self.converted_dir = self.output_dir / "converted"
        self.logs_dir = self.output_dir / "logs"
        for d in [self.original_dir, self.converted_dir, self.logs_dir]:
            d.mkdir(parents=True, exist_ok=True)

    def is_scanned_pdf(self, pdf_path: str) -> bool:
        """Check if PDF is likely scanned by analyzing text extraction"""
        if not PYPDF2_AVAILABLE:
            return False

        try:
            reader = PdfReader(pdf_path)
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() or ""

            if len(text_content.strip()) < 100:
                return True

            text_ratio = len(text_content) / (len(reader.pages) * 1000)
            return text_ratio < 0.1
        except Exception as e:
            logger.error(f"Error checking if PDF is scanned: {e}")
            return False

    def convert_pdf_to_word(self, pdf_path: str) -> Tuple[bool, str, str]:
        """
        Convert PDF to Word document

        Returns:
            Tuple of (success: bool, output_path: str, message: str)
        """
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            return False, "", f"PDF file not found: {pdf_path}"

        is_scanned = self.is_scanned_pdf(str(pdf_path))
        logger.info(f"PDF '{pdf_path.name}' is {'scanned' if is_scanned else 'text-based'}")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{pdf_path.stem}_{timestamp}.docx"
        output_path = self.converted_dir / output_filename

        original_copy = self.original_dir / f"{pdf_path.stem}_{timestamp}.pdf"
        shutil.copy2(pdf_path, original_copy)

        if pdf2docx_AVAILABLE:
            try:
                cv = Converter(str(pdf_path))
                cv.convert(str(output_path), start=0, end=None)
                cv.close()
                logger.info(f"Successfully converted '{pdf_path.name}' to '{output_path.name}'")
                return True, str(output_path), f"Successfully converted (pdf2docx)"
            except Exception as e:
                logger.error(f"pdf2docx conversion failed: {e}")
                return False, "", f"Conversion failed: {e}"

        elif PYPDF2_AVAILABLE:
            return self._fallback_convert(str(pdf_path), output_path)
        else:
            return False, "", "No PDF conversion library available. Please install pdf2docx or PyPDF2."

    def _fallback_convert(self, pdf_path: str, output_path: Path) -> Tuple[bool, str, str]:
        """Fallback text extraction using PyPDF2"""
        try:
            from docx import Document
            reader = PdfReader(pdf_path)
            doc = Document()
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    doc.add_heading(f"Page {i+1}", level=1)
                    doc.add_paragraph(text)
            doc.save(output_path)
            logger.info(f"Fallback conversion completed: {output_path.name}")
            return True, str(output_path), "Converted (PyPDF2 fallback - text only)"
        except Exception as e:
            logger.error(f"Fallback conversion failed: {e}")
            return False, "", f"Fallback conversion failed: {e}"

    def batch_convert(self, pdf_dir: str) -> List[dict]:
        """Convert all PDFs in a directory"""
        pdf_dir = Path(pdf_dir)
        results = []
        for pdf_file in pdf_dir.glob("*.pdf"):
            success, output_path, message = self.convert_pdf_to_word(str(pdf_file))
            results.append({
                "file": pdf_file.name,
                "success": success,
                "output": output_path,
                "message": message
            })
        return results

    def get_output_dirs(self) -> dict:
        """Return paths to all output directories"""
        return {
            "original": str(self.original_dir),
            "converted": str(self.converted_dir),
            "logs": str(self.logs_dir)
        }


class LogManager:
    """Manage operation logs"""

    def __init__(self, logs_dir: str = "output/logs"):
        self.logs_dir = Path(logs_dir)
        self.logs_dir.mkdir(parents=True, exist_ok=True)

    def log_operation(self, operation: str, filename: str, status: str, details: str = ""):
        """Log an operation with timestamp"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_file = self.logs_dir / f"operation_{datetime.now().strftime('%Y%m%d')}.log"
        log_entry = f"[{timestamp}] [{operation.upper()}] {filename} - {status}"
        if details:
            log_entry += f" - {details}"
        log_entry += "\n"
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(log_entry)
        logger.info(log_entry.strip())

    def get_recent_logs(self, limit: int = 50) -> List[str]:
        """Get recent log entries"""
        log_files = sorted(self.logs_dir.glob("operation_*.log"), reverse=True)
        logs = []
        for log_file in log_files[:5]:
            with open(log_file, "r", encoding="utf-8") as f:
                logs.extend(f.readlines())
        return logs[-limit:]


def main():
    import tempfile
    converter = PDFToWordConverter(output_dir="output")
    log_manager = LogManager(logs_dir="output/logs")

    test_pdf = None
    if test_pdf and Path(test_pdf).exists():
        success, output_path, message = converter.convert_pdf_to_word(test_pdf)
        log_manager.log_operation("convert", test_pdf, "success" if success else "failed", message)
        print(f"Result: {message}")
        if success:
            print(f"Output: {output_path}")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()
